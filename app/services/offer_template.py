from __future__ import annotations

import html
import io
import mimetypes
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any
from xml.etree import ElementTree as ET

import fitz
import httpx
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu
from docx.text.paragraph import Paragraph

from app.models.rfq import Rfq

WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
WORD_TEXT_TAG = f"{{{WORD_NAMESPACE}}}t"
WORD_PARAGRAPH_TAG = f"{{{WORD_NAMESPACE}}}p"
XML_SPACE_ATTRIBUTE = "{http://www.w3.org/XML/1998/namespace}space"
PLACEHOLDER_PATTERN = re.compile(r"\{\{\s*([A-Za-z0-9_]+)\s*\}\}")
XMLNS_DECLARATION_PATTERN = re.compile(r"""xmlns:([A-Za-z0-9_]+)=(['"])[^'"]+\2""")
IGNORABLE_ATTRIBUTE_PATTERN = re.compile(
    r"""([A-Za-z0-9_]+:Ignorable)=(['"])([^'"]*)\2"""
)
BLANK_WHEN_MISSING_PLACEHOLDERS = {
    "created_by_name",
    "created_by_phone",
    "created_by_email",
}
WEEKDAY_NAMES = (
    "Monday",
    "Tuesday",
    "Wednesday",
    "Thursday",
    "Friday",
    "Saturday",
    "Sunday",
)
MONTH_NAMES = (
    "",
    "January",
    "February",
    "March",
    "April",
    "May",
    "June",
    "July",
    "August",
    "September",
    "October",
    "November",
    "December",
)

APP_ROOT = Path(__file__).resolve().parents[1]
OFFER_TEMPLATE_PATH = APP_ROOT / "assets" / "offer_preparation_template.docx"
BACKEND_ROOT = APP_ROOT.parent
REPO_ROOT = BACKEND_ROOT.parent
DOCX_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff"}
OFFER_PREPARATION_DATA_KEY = "offer_preparation_data"
DEFAULT_OFFER_SUBJECT = "commercial offer AVOCarbon TUNISIA-SAME"
DEFAULT_OFFER_VALIDITY = "1 month from date of issue."
DRAWING_FILE_ROLE_HINTS = {"DRAWING", "RFQ_DRAWING", "ANNEX_DRAWING"}
DRAWING_NAME_HINTS = ("drawing", "plan", "schematic", "schema", "blueprint", "dwg")
REFERENCE_PICTURE_NAME_HINTS = ("reference", "picture", "photo", "image", "product")
REFERENCE_PICTURE_WIDTH_RATIO = 0.78
REFERENCE_PICTURE_BEFORE_ROLE_HINTS = {
    "REFERENCE_PICTURE_BEFORE",
    "REFERENCE_BEFORE",
    "PRODUCT_PICTURE_BEFORE",
}
REFERENCE_PICTURE_AFTER_ROLE_HINTS = {
    "REFERENCE_PICTURE_AFTER",
    "REFERENCE_AFTER",
    "PRODUCT_PICTURE_AFTER",
}
REFERENCE_PICTURE_BEFORE_NAME_HINTS = (
    "before-product-picture",
    "before_reference",
    "reference-before",
    "photo-before",
    "before-photo",
)
REFERENCE_PICTURE_AFTER_NAME_HINTS = (
    "after-product-picture",
    "after_reference",
    "reference-after",
    "photo-after",
    "after-photo",
)


def build_offer_preparation_filename(rfq: Rfq) -> str:
    return OFFER_TEMPLATE_PATH.name


def render_offer_preparation_docx(
    rfq: Rfq,
    creator_profile: dict[str, Any] | None = None,
    offer_data: dict[str, Any] | None = None,
) -> bytes:
    if not OFFER_TEMPLATE_PATH.exists():
        raise RuntimeError(f"Offer template not found: {OFFER_TEMPLATE_PATH}")

    rfq_data = dict(rfq.rfq_data or {})
    context = build_offer_preparation_context(
        rfq,
        creator_profile=creator_profile,
        offer_data=offer_data,
    )
    with OFFER_TEMPLATE_PATH.open("rb") as template_file:
        template_bytes = template_file.read()

    rendered_docx = _render_docx_template(template_bytes, context)
    return _inject_offer_previews_into_docx(rendered_docx, rfq_data, context)


def render_offer_preparation_preview_html(
    rfq: Rfq,
    creator_profile: dict[str, Any] | None = None,
    offer_data: dict[str, Any] | None = None,
) -> str:
    filled_docx = render_offer_preparation_docx(
        rfq,
        creator_profile=creator_profile,
        offer_data=offer_data,
    )
    paragraphs = _extract_docx_paragraphs(filled_docx)
    paragraph_html = "".join(_render_preview_paragraph(text) for text in paragraphs)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
  <meta charset="utf-8" />
  <title>Offer preparation preview</title>
  <style>
    :root {{
      color-scheme: light;
      --page-bg: #eef2f7;
      --paper: #ffffff;
      --ink: #111827;
      --muted: #6b7280;
      --line: #d1d5db;
    }}
    * {{
      box-sizing: border-box;
    }}
    body {{
      margin: 0;
      background: var(--page-bg);
      color: var(--ink);
      font-family: "Times New Roman", Times, serif;
      line-height: 1.5;
      padding: 24px 16px 40px;
    }}
    .paper {{
      width: min(100%, 900px);
      margin: 0 auto;
      background: var(--paper);
      border: 1px solid var(--line);
      box-shadow: 0 22px 50px rgba(15, 23, 42, 0.08);
      padding: 40px 48px 52px;
    }}
    .content {{
      width: 100%;
    }}
    .content p,
    .content h2,
    .content h3 {{
      margin: 0 0 12px;
      white-space: pre-wrap;
      word-break: break-word;
    }}
    .content h2 {{
      margin-top: 20px;
      padding-top: 10px;
      border-top: 1px solid var(--line);
      font-size: 18px;
      font-weight: 700;
    }}
    .content h3 {{
      margin-top: 12px;
      font-size: 16px;
      font-weight: 700;
    }}
    .content .label {{
      font-weight: 700;
    }}
    .content .muted {{
      color: var(--muted);
    }}
    .content .highlight {{
      padding-left: 18px;
    }}
    @media (max-width: 720px) {{
      body {{
        padding: 12px;
      }}
      .paper {{
        padding: 28px 22px 36px;
      }}
    }}
  </style>
</head>
<body>
  <div class="paper">
    <div class="content">{paragraph_html}</div>
  </div>
</body>
</html>
"""


def build_offer_preparation_context(
    rfq: Rfq,
    creator_profile: dict[str, Any] | None = None,
    offer_data: dict[str, Any] | None = None,
) -> dict[str, str]:
    rfq_data = dict(rfq.rfq_data or {})
    data = _build_effective_offer_template_data(rfq_data, offer_data=offer_data)
    creator_data = dict(creator_profile or {})
    latest_rfq_file = _select_annex_drawing_file(_get_normalized_rfq_files(rfq_data))
    return {
        "created_at": _format_offer_created_at(rfq.created_at),
        "our_reference": _build_offer_reference(rfq, offer_data=offer_data),
        "rfq_file_name": _stringify_value(latest_rfq_file.get("name")),
        "your_reference": _build_your_reference(data),
        "copies": _pick_first_value(data, ("copies", "copy_recipients", "cc_recipients")),
        "subject": _build_offer_subject(data),
        "created_by_name": _stringify_value(
            creator_data.get("created_by_name")
            or creator_data.get("full_name")
            or rfq.created_by_email
        ),
        "created_by_phone": _stringify_value(
            creator_data.get("created_by_phone") or creator_data.get("phone")
        ),
        "created_by_email": _stringify_value(
            creator_data.get("created_by_email") or rfq.created_by_email
        ),
        "customer_name": _pick_first_value(data, ("customer_name", "customer", "client")),
        "contact_name": _pick_first_value(
            data,
            ("contact_name", "contact_first_name", "contactName"),
        ),
        "contact_phone": _pick_first_value(data, ("contact_phone", "contactPhone")),
        "contact_email": _pick_first_value(data, ("contact_email", "contactEmail")),
        "product_name": _pick_first_value(data, ("product_name", "productName")),
        "project_name": _pick_first_value(data, ("project_name", "projectName")),
        "customer_pn": _pick_first_value(data, ("customer_pn", "customerPn")),
        "revision_level": _pick_first_value(
            data,
            ("revision_level", "revisionLevel"),
        ),
        "sop_year": _pick_first_value(data, ("sop_year", "sop")),
        "validation_batch": _pick_first_value(
            data,
            ("validation_batch", "validationBatch"),
        ),
        "annual_volume": _pick_first_value(
            data,
            ("annual_volume", "qty_per_year", "qtyPerYear"),
        ),
        "pilot_quantity": _pick_first_value(
            data,
            ("pilot_quantity", "pilotQuantity"),
        ),
        "pilot_unit_price": _pick_first_value(
            data,
            ("pilot_unit_price", "pilotUnitPrice"),
        ),
        "material_balance_moq": _pick_first_value(
            data,
            ("material_balance_moq", "materialBalanceMoq", "moq"),
        ),
        "serial_unit_price": _build_serial_unit_price_display(data),
        "type_of_packaging": _pick_first_value(
            data,
            ("type_of_packaging", "typeOfPackaging"),
        ),
        "lead_time_deliveries": _pick_first_value(
            data,
            ("lead_time_deliveries", "leadTimeDeliveries"),
        ),
        "expected_delivery_conditions": _pick_first_value(
            data,
            ("expected_delivery_conditions", "expectedDeliveryConditions"),
        ),
        "expected_payment_terms": _pick_first_value(
            data,
            ("expected_payment_terms", "expectedPaymentTerms"),
        ),
        "inventory_commitment": _pick_first_value(
            data,
            ("inventory_commitment", "inventoryCommitment"),
        ),
        "offer_validity": _build_offer_validity(data),
        "target_price_display": _format_target_price_display(data),
    }


def _render_docx_template(template_bytes: bytes, context: dict[str, str]) -> bytes:
    source_buffer = io.BytesIO(template_bytes)
    output_buffer = io.BytesIO()

    with zipfile.ZipFile(source_buffer, "r") as source_zip, zipfile.ZipFile(
        output_buffer, "w", zipfile.ZIP_DEFLATED
    ) as output_zip:
        for item in source_zip.infolist():
            file_bytes = source_zip.read(item.filename)
            if item.filename.startswith("word/") and item.filename.endswith(".xml"):
                file_bytes = _replace_placeholders_in_xml(file_bytes, context)
            output_zip.writestr(item, file_bytes)

    return output_buffer.getvalue()


def _replace_placeholders_in_xml(xml_bytes: bytes, context: dict[str, str]) -> bytes:
    root = ET.fromstring(xml_bytes)
    _replace_placeholder_text_nodes(root, context)
    _apply_offer_line_overrides(root, context)
    rendered_bytes = ET.tostring(root, encoding="utf-8", xml_declaration=True)
    return _sanitize_ignorable_prefixes(rendered_bytes)


def _replace_placeholder_text_nodes(root: ET.Element, context: dict[str, str]) -> None:
    for paragraph in [node for node in root.iter() if node.tag == WORD_PARAGRAPH_TAG]:
        text_nodes = [node for node in paragraph.iter() if node.tag == WORD_TEXT_TAG]
        if not text_nodes:
            continue

        full_text = "".join(node.text or "" for node in text_nodes)
        if "{{" not in full_text:
            continue

        _replace_placeholders_across_text_nodes(text_nodes, context)


def _replace_placeholders_across_text_nodes(
    text_nodes: list[ET.Element],
    context: dict[str, str],
) -> None:
    original_texts = [node.text or "" for node in text_nodes]
    full_text = "".join(original_texts)
    matches = list(PLACEHOLDER_PATTERN.finditer(full_text))
    if not matches:
        return

    replacements_by_start = {
        match.start(): (
            match.end(),
            _placeholder_replacement_text(match.group(1), context),
        )
        for match in matches
    }

    global_offset = 0
    skip_until = -1

    for node, original_text in zip(text_nodes, original_texts):
        rebuilt: list[str] = []
        for local_index, char in enumerate(original_text):
            global_index = global_offset + local_index
            if global_index < skip_until:
                continue

            replacement = replacements_by_start.get(global_index)
            if replacement is not None:
                replacement_end, replacement_text = replacement
                rebuilt.append(replacement_text)
                skip_until = replacement_end
                continue

            rebuilt.append(char)

        _set_text_node_value(node, "".join(rebuilt))
        global_offset += len(original_text)


def _apply_offer_line_overrides(root: ET.Element, context: dict[str, str]) -> None:
    unit_price_occurrence = 0
    for paragraph in [node for node in root.iter() if node.tag == WORD_PARAGRAPH_TAG]:
        text_nodes = [node for node in paragraph.iter() if node.tag == WORD_TEXT_TAG]
        if not text_nodes:
            continue

        original_texts = [node.text or "" for node in text_nodes]
        normalized_text = _normalize_preview_text("".join(original_texts))
        if not normalized_text:
            continue

        replacement_text = ""
        if normalized_text.startswith("Our reference:") and _has_meaningful_value(
            context.get("our_reference")
        ):
            replacement_text = f"Our reference: {context['our_reference']}"
        elif normalized_text == "Cell :" and _has_meaningful_value(
            context.get("created_by_phone")
        ):
            replacement_text = f"Cell : {context['created_by_phone']}"
        elif normalized_text == "E-mail :" and _has_meaningful_value(
            context.get("created_by_email")
        ):
            replacement_text = f"E-mail : {context['created_by_email']}"
        elif normalized_text == "Unit price:":
            unit_price_occurrence += 1
            if (
                unit_price_occurrence == 2
                and _has_meaningful_value(context.get("serial_unit_price"))
            ):
                replacement_text = f"Unit price: {context['serial_unit_price']}"

        if replacement_text:
            _replace_entire_text_node_sequence(text_nodes, replacement_text)


def _replace_entire_text_node_sequence(
    text_nodes: list[ET.Element],
    replacement_text: str,
) -> None:
    for index, node in enumerate(text_nodes):
        _set_text_node_value(node, replacement_text if index == 0 else "")


def _set_text_node_value(node: ET.Element, value: str) -> None:
    node.text = value
    if value and (value.startswith(" ") or value.endswith(" ")):
        node.set(XML_SPACE_ATTRIBUTE, "preserve")
    else:
        node.attrib.pop(XML_SPACE_ATTRIBUTE, None)


def _placeholder_replacement_text(name: str, context: dict[str, str]) -> str:
    value = _stringify_value(context.get(name))
    if value:
        return value
    if name in BLANK_WHEN_MISSING_PLACEHOLDERS:
        return ""
    return "-"


def _format_offer_created_at(value: datetime | None) -> str:
    if value is None:
        value = datetime.now(timezone.utc)
    return (
        f"{WEEKDAY_NAMES[value.weekday()]}, "
        f"{MONTH_NAMES[value.month]} {value.day}, {value.year}"
    )


def _build_offer_reference(
    rfq: Rfq,
    offer_data: dict[str, Any] | None = None,
) -> str:
    data = _build_effective_offer_template_data(
        dict(rfq.rfq_data or {}),
        offer_data=offer_data,
    )
    created_at = rfq.created_at or datetime.now(timezone.utc)
    parts = [
        created_at.strftime("%Y%m%d"),
        _pick_first_value(data, ("project_name", "projectName")),
        _pick_first_value(data, ("product_name", "productName")),
        _pick_first_value(data, ("customer_name", "customer", "client")),
    ]
    cleaned_parts = [_normalize_reference_part(part) for part in parts if part]
    return "-".join(part for part in cleaned_parts if part)


def _build_your_reference(data: dict[str, Any]) -> str:
    explicit_reference = _pick_first_value(
        data,
        ("your_reference", "yourReference"),
    )
    if explicit_reference:
        return explicit_reference

    project_name = _pick_first_value(data, ("project_name", "projectName"))
    product_name = _pick_first_value(data, ("product_name", "productName"))
    reference_parts = [
        _stringify_value(part)
        for part in (project_name, product_name)
        if _has_meaningful_value(part)
    ]
    if not reference_parts:
        return ""
    return f"RFQ/{' '.join(reference_parts)}"


def _build_offer_subject(data: dict[str, Any]) -> str:
    return _pick_first_value(data, ("subject", "offer_subject", "offerSubject")) or DEFAULT_OFFER_SUBJECT


def _build_serial_unit_price_display(data: dict[str, Any]) -> str:
    explicit_price = _pick_first_value(
        data,
        ("serial_unit_price", "serialUnitPrice"),
    )
    if explicit_price:
        return explicit_price
    return _format_target_price_display(data)


def _build_offer_validity(data: dict[str, Any]) -> str:
    return _pick_first_value(data, ("offer_validity", "offerValidity")) or DEFAULT_OFFER_VALIDITY


def _inject_offer_previews_into_docx(
    docx_bytes: bytes,
    rfq_data: dict[str, Any],
    context: dict[str, str],
) -> bytes:
    document = Document(io.BytesIO(docx_bytes))
    _apply_offer_document_overrides(document, context)

    normalized_files = _get_normalized_rfq_files(rfq_data)
    annex_file = _select_annex_drawing_file(normalized_files)
    reference_picture_groups = _select_reference_picture_groups(normalized_files, annex_file)

    _replace_annex_drawing_preview(document, annex_file)
    _replace_reference_picture_previews(document, reference_picture_groups)

    output_buffer = io.BytesIO()
    document.save(output_buffer)
    return output_buffer.getvalue()


def _apply_offer_document_overrides(
    document: Document,
    context: dict[str, str],
) -> None:
    _replace_paragraph_text_by_prefix(
        document,
        "Our reference:",
        f"Our reference: {context['our_reference']}",
    )

    if _has_meaningful_value(context.get("your_reference")):
        _replace_paragraph_text_by_prefix(
            document,
            "Your reference:",
            f"Your reference: {context['your_reference']}",
        )

    if _has_meaningful_value(context.get("copies")):
        _replace_paragraph_text_by_prefix(
            document,
            "Copies :",
            f"Copies : {context['copies']}",
        )

    if _has_meaningful_value(context.get("subject")):
        _replace_paragraph_text_by_prefix(
            document,
            "SUBJECT :",
            f"SUBJECT : {context['subject']}",
        )

    if _has_meaningful_value(context.get("validation_batch")):
        scope_paragraph = _find_paragraph_by_prefix(document, "Product:")
        if scope_paragraph is not None:
            _replace_line_in_multiline_paragraph(
                scope_paragraph,
                "Validation batch:",
                f"Validation batch: {context['validation_batch']}",
            )

    pilot_heading_index = _find_paragraph_index_by_prefix(
        document,
        "A. Validation batch (Pilot)",
    )
    if pilot_heading_index >= 0:
        _replace_labeled_paragraph_after_index(
            document,
            pilot_heading_index,
            "Quantity:",
            context.get("pilot_quantity") or "",
        )
        _replace_labeled_paragraph_after_index(
            document,
            pilot_heading_index,
            "Unit price:",
            context.get("pilot_unit_price") or "",
        )
        _replace_labeled_paragraph_after_index(
            document,
            pilot_heading_index,
            "Material balance (MOQ):",
            context.get("material_balance_moq") or "",
        )

    serial_heading_index = _find_paragraph_index_by_prefix(document, "B. Serial production")
    if serial_heading_index >= 0:
        _replace_labeled_paragraph_after_index(
            document,
            serial_heading_index,
            "Unit price:",
            context.get("serial_unit_price") or "",
        )

    if _has_meaningful_value(context.get("lead_time_deliveries")):
        _replace_following_content_paragraph(
            document,
            "4) Lead time & deliveries",
            context["lead_time_deliveries"],
        )

    if _has_meaningful_value(context.get("inventory_commitment")):
        _replace_following_content_paragraph(
            document,
            "Inventory commitment",
            context["inventory_commitment"],
        )

    if _has_meaningful_value(context.get("offer_validity")):
        _replace_paragraph_text_by_prefix(
            document,
            "Offer validity:",
            f"Offer validity: {context['offer_validity']}",
        )


def _replace_annex_drawing_preview(
    document: Document,
    annex_file: dict[str, str] | None,
) -> None:
    annex_file = annex_file or {}
    annex_heading_index = _find_paragraph_index_by_prefix(document, "Annex1 : Drawing")
    if annex_heading_index < 0:
        return

    target_paragraph = _find_or_create_annex_content_paragraph(document, annex_heading_index)
    _clear_docx_paragraph(target_paragraph)

    rfq_file_name = _stringify_value(annex_file.get("name"))
    if not rfq_file_name:
        return

    preview_image = _build_rfq_file_preview_image(annex_file)
    if preview_image is not None:
        _add_image_to_paragraph(
            target_paragraph,
            preview_image,
            width=_get_section_available_width(document),
        )
        return

    target_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    target_paragraph.add_run(f"RFQ file: {rfq_file_name}")


def _replace_reference_picture_previews(
    document: Document,
    reference_picture_groups: dict[str, list[dict[str, str]]],
) -> None:
    before_files = list(reference_picture_groups.get("before") or [])
    after_files = list(reference_picture_groups.get("after") or [])
    if not before_files and not after_files:
        return

    heading_index = _find_paragraph_index_by_prefix(document, "Product picture for reference")
    if heading_index < 0:
        return

    heading_paragraph = document.paragraphs[heading_index]
    available_width = _get_section_available_width(document)
    target_width = Emu(int(available_width * REFERENCE_PICTURE_WIDTH_RATIO))

    before_payloads = _build_reference_picture_payloads(before_files)
    after_payloads = _build_reference_picture_payloads(after_files)

    if before_payloads:
        target_paragraph = _insert_paragraph_before(heading_paragraph)
        _render_reference_picture_sequence(target_paragraph, before_payloads, target_width)

    if after_payloads:
        target_paragraph = _insert_paragraph_after(heading_paragraph)
        _render_reference_picture_sequence(target_paragraph, after_payloads, target_width)


def _add_image_to_paragraph(
    paragraph: Paragraph,
    image_bytes: bytes,
    *,
    width: Emu,
) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=width)


def _build_reference_picture_payloads(
    reference_picture_files: list[dict[str, str]],
) -> list[bytes]:
    image_payloads: list[bytes] = []
    for file_entry in reference_picture_files:
        image_bytes = _build_rfq_file_preview_image(file_entry)
        if image_bytes is not None:
            image_payloads.append(image_bytes)
    return image_payloads


def _render_reference_picture_sequence(
    target_paragraph: Paragraph,
    image_payloads: list[bytes],
    width: Emu,
) -> None:
    current_paragraph = target_paragraph
    for index, image_bytes in enumerate(image_payloads):
        if index > 0:
            current_paragraph = _insert_paragraph_after(current_paragraph)
        _clear_docx_paragraph(current_paragraph)
        _add_image_to_paragraph(current_paragraph, image_bytes, width=width)


def _replace_paragraph_text_by_prefix(
    document: Document,
    prefix: str,
    replacement_text: str,
) -> None:
    paragraph = _find_paragraph_by_prefix(document, prefix)
    if paragraph is None or not replacement_text:
        return
    paragraph.text = replacement_text


def _replace_labeled_paragraph_after_index(
    document: Document,
    start_index: int,
    label_prefix: str,
    value: str,
) -> None:
    if not _has_meaningful_value(value):
        return
    paragraph = _find_paragraph_after_index_by_prefix(document, start_index, label_prefix)
    if paragraph is None:
        return
    paragraph.text = f"{label_prefix} {value}"


def _replace_following_content_paragraph(
    document: Document,
    heading_prefix: str,
    replacement_text: str,
) -> None:
    heading_index = _find_paragraph_index_by_prefix(document, heading_prefix)
    if heading_index < 0 or not replacement_text:
        return
    paragraph = _find_or_create_following_content_paragraph(document, heading_index)
    paragraph.text = replacement_text


def _replace_line_in_multiline_paragraph(
    paragraph: Paragraph,
    line_prefix: str,
    replacement_text: str,
) -> None:
    lines = paragraph.text.splitlines() or [paragraph.text]
    updated_lines: list[str] = []
    line_was_replaced = False

    for line in lines:
        if _normalize_preview_text(line).startswith(line_prefix):
            updated_lines.append(replacement_text)
            line_was_replaced = True
        else:
            updated_lines.append(line)

    if not line_was_replaced:
        updated_lines.append(replacement_text)

    paragraph.text = "\n".join(updated_lines)


def _find_paragraph_by_prefix(document: Document, prefix: str) -> Paragraph | None:
    index = _find_paragraph_index_by_prefix(document, prefix)
    if index < 0:
        return None
    return document.paragraphs[index]


def _find_paragraph_index_by_prefix(document: Document, prefix: str) -> int:
    for index, paragraph in enumerate(document.paragraphs):
        if _normalize_preview_text(paragraph.text).startswith(prefix):
            return index
    return -1


def _find_paragraph_after_index_by_prefix(
    document: Document,
    start_index: int,
    prefix: str,
) -> Paragraph | None:
    for paragraph in document.paragraphs[start_index + 1 :]:
        if _normalize_preview_text(paragraph.text).startswith(prefix):
            return paragraph
    return None


def _find_or_create_following_content_paragraph(
    document: Document,
    heading_index: int,
) -> Paragraph:
    for paragraph in document.paragraphs[heading_index + 1 :]:
        if _normalize_preview_text(paragraph.text):
            break
        return paragraph

    return _insert_paragraph_after(document.paragraphs[heading_index])


def _build_rfq_file_preview_image(latest_rfq_file: dict[str, str]) -> bytes | None:
    file_bytes = _load_rfq_file_bytes(latest_rfq_file)
    if not file_bytes:
        return None

    content_type = _detect_rfq_file_content_type(latest_rfq_file)
    extension = _get_rfq_file_extension(latest_rfq_file)

    if content_type == "application/pdf" or extension == ".pdf":
        return _render_pdf_preview_image(file_bytes)

    if extension in DOCX_IMAGE_EXTENSIONS or content_type.startswith("image/"):
        return file_bytes

    return None


def _load_rfq_file_bytes(latest_rfq_file: dict[str, str]) -> bytes | None:
    for source in (
        latest_rfq_file.get("download_url"),
        latest_rfq_file.get("url"),
        latest_rfq_file.get("path"),
        latest_rfq_file.get("blob_url"),
    ):
        normalized_source = _stringify_value(source)
        if not normalized_source:
            continue

        if normalized_source.startswith(("http://", "https://")):
            try:
                response = httpx.get(normalized_source, follow_redirects=True, timeout=20.0)
                response.raise_for_status()
                return response.content
            except httpx.HTTPError:
                continue

        legacy_local_file = _resolve_legacy_rfq_local_file(normalized_source)
        if legacy_local_file is not None:
            try:
                return legacy_local_file.read_bytes()
            except OSError:
                continue

        candidate_path = Path(normalized_source)
        if candidate_path.exists() and candidate_path.is_file():
            try:
                return candidate_path.read_bytes()
            except OSError:
                continue

    return None


def _resolve_legacy_rfq_local_file(source: str) -> Path | None:
    if "/api/rfq/download/" not in source:
        return None

    stored_name = source.rsplit("/", 1)[-1]
    for base_root in (BACKEND_ROOT, REPO_ROOT):
        candidate = base_root / "uploads" / stored_name
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def _render_pdf_preview_image(pdf_bytes: bytes) -> bytes | None:
    try:
        with fitz.open(stream=pdf_bytes, filetype="pdf") as pdf_document:
            if pdf_document.page_count == 0:
                return None
            page = pdf_document.load_page(0)
            pixmap = page.get_pixmap(matrix=fitz.Matrix(1.6, 1.6), alpha=False)
            return pixmap.tobytes("png")
    except Exception:
        return None


def _find_or_create_annex_content_paragraph(
    document: Document,
    annex_heading_index: int,
) -> Paragraph:
    for paragraph in document.paragraphs[annex_heading_index + 1 :]:
        if _normalize_preview_text(paragraph.text):
            break
        return paragraph

    return _insert_paragraph_after(document.paragraphs[annex_heading_index])


def _insert_paragraph_after(paragraph: Paragraph) -> Paragraph:
    new_paragraph_xml = OxmlElement("w:p")
    paragraph._p.addnext(new_paragraph_xml)
    return Paragraph(new_paragraph_xml, paragraph._parent)


def _insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_paragraph_xml = OxmlElement("w:p")
    paragraph._p.addprevious(new_paragraph_xml)
    return Paragraph(new_paragraph_xml, paragraph._parent)


def _clear_docx_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def _get_section_available_width(document: Document) -> Emu:
    section = document.sections[0]
    return Emu(int(section.page_width - section.left_margin - section.right_margin))


def _sanitize_ignorable_prefixes(xml_bytes: bytes) -> bytes:
    xml_text = xml_bytes.decode("utf-8")
    root_start_index = xml_text.find("<")
    if root_start_index == -1:
        return xml_bytes

    if xml_text.startswith("<?xml"):
        root_start_index = xml_text.find("<", xml_text.find("?>") + 2)
        if root_start_index == -1:
            return xml_bytes

    root_end_index = xml_text.find(">", root_start_index)
    if root_end_index == -1:
        return xml_bytes

    root_open_tag = xml_text[root_start_index : root_end_index + 1]
    ignorable_match = IGNORABLE_ATTRIBUTE_PATTERN.search(root_open_tag)
    if ignorable_match is None:
        return xml_bytes

    declared_prefixes = {
        match.group(1) for match in XMLNS_DECLARATION_PATTERN.finditer(root_open_tag)
    }
    ignorable_prefixes = ignorable_match.group(3).split()
    filtered_prefixes = [
        prefix for prefix in ignorable_prefixes if prefix in declared_prefixes
    ]

    if filtered_prefixes == ignorable_prefixes:
        return xml_bytes

    if filtered_prefixes:
        replacement = (
            f'{ignorable_match.group(1)}={ignorable_match.group(2)}'
            f'{" ".join(filtered_prefixes)}'
            f"{ignorable_match.group(2)}"
        )
    else:
        replacement = ""

    sanitized_root_open_tag = (
        root_open_tag[: ignorable_match.start()]
        + replacement
        + root_open_tag[ignorable_match.end() :]
    )
    sanitized_root_open_tag = re.sub(r"\s{2,}", " ", sanitized_root_open_tag)
    sanitized_xml_text = (
        xml_text[:root_start_index]
        + sanitized_root_open_tag
        + xml_text[root_end_index + 1 :]
    )
    return sanitized_xml_text.encode("utf-8")


def _normalize_reference_part(value: str) -> str:
    normalized = re.sub(r"\s+", " ", _stringify_value(value))
    return normalized.strip(" -")


def _get_latest_rfq_file(data: dict[str, Any]) -> dict[str, str]:
    normalized_files = _get_normalized_rfq_files(data)
    return normalized_files[0] if normalized_files else {}


def _get_normalized_rfq_files(data: dict[str, Any]) -> list[dict[str, str]]:
    raw_files = data.get("rfq_files")
    if not isinstance(raw_files, list):
        return []

    normalized_files: list[dict[str, str]] = []
    for entry in raw_files:
        if isinstance(entry, str):
            normalized_files.append(
                {
                    "name": Path(entry).name,
                    "uploaded_at": "",
                    "path": entry,
                    "url": entry,
                    "download_url": entry,
                    "blob_url": "",
                    "content_type": _guess_content_type(entry),
                    "file_role": "",
                    "blob_name": "",
                    "id": "",
                }
            )
            continue
        if not isinstance(entry, dict):
            continue
        normalized_files.append(
            {
                "name": _stringify_value(
                    entry.get("name")
                    or entry.get("filename")
                    or entry.get("original_name")
                    or entry.get("file_name")
                ),
                "uploaded_at": _stringify_value(
                    entry.get("uploaded_at")
                    or entry.get("updated_at")
                    or entry.get("last_modified")
                ),
                "path": _stringify_value(entry.get("path")),
                "url": _stringify_value(entry.get("url")),
                "download_url": _stringify_value(entry.get("download_url")),
                "blob_url": _stringify_value(entry.get("blob_url")),
                "content_type": _stringify_value(entry.get("content_type")),
                "file_role": _stringify_value(entry.get("file_role")),
                "blob_name": _stringify_value(entry.get("blob_name")),
                "id": _stringify_value(entry.get("id") or entry.get("file_id")),
            }
        )

    normalized_files = [
        file_entry for file_entry in normalized_files if _has_meaningful_value(file_entry.get("name"))
    ]
    if not normalized_files:
        return []

    normalized_files.sort(
        key=lambda file_entry: _parse_sortable_timestamp(file_entry.get("uploaded_at")),
        reverse=True,
    )
    return normalized_files


def _select_annex_drawing_file(files: list[dict[str, str]]) -> dict[str, str]:
    if not files:
        return {}

    selectors = (
        lambda file_entry: _normalize_file_role(file_entry) in DRAWING_FILE_ROLE_HINTS,
        lambda file_entry: _is_pdf_file(file_entry) and _has_drawing_name_hint(file_entry),
        lambda file_entry: _is_pdf_file(file_entry),
        lambda file_entry: _has_drawing_name_hint(file_entry) and _is_previewable_file(file_entry),
        lambda file_entry: _is_previewable_file(file_entry),
        lambda _file_entry: True,
    )

    for selector in selectors:
        for file_entry in files:
            if selector(file_entry):
                return file_entry

    return {}


def _select_reference_picture_groups(
    files: list[dict[str, str]],
    annex_file: dict[str, str] | None = None,
) -> dict[str, list[dict[str, str]]]:
    annex_file = annex_file or {}
    candidate_images = [
        file_entry
        for file_entry in files
        if _is_image_file(file_entry) and not _is_same_rfq_file(file_entry, annex_file)
    ]
    if not candidate_images:
        return {"before": [], "after": []}

    before_images: list[dict[str, str]] = []
    after_images: list[dict[str, str]] = []
    fallback_images: list[dict[str, str]] = []

    for file_entry in candidate_images:
        if _is_reference_picture_before_file(file_entry):
            before_images.append(file_entry)
        elif _is_reference_picture_after_file(file_entry):
            after_images.append(file_entry)
        elif _has_reference_picture_name_hint(file_entry):
            fallback_images.append(file_entry)

    if not before_images and not after_images:
        before_images = fallback_images or candidate_images
    else:
        remaining_images = [
            file_entry
            for file_entry in candidate_images
            if file_entry not in before_images and file_entry not in after_images
        ]
        before_images.extend(
            file_entry
            for file_entry in remaining_images
            if file_entry not in before_images
        )

    return {
        "before": sorted(
            before_images,
            key=lambda file_entry: _parse_sortable_timestamp(file_entry.get("uploaded_at")),
        ),
        "after": sorted(
            after_images,
            key=lambda file_entry: _parse_sortable_timestamp(file_entry.get("uploaded_at")),
        ),
    }


def _parse_sortable_timestamp(value: str | None) -> float:
    normalized_value = _stringify_value(value)
    if not normalized_value:
        return 0.0
    try:
        return datetime.fromisoformat(normalized_value.replace("Z", "+00:00")).timestamp()
    except ValueError:
        return 0.0


def _detect_rfq_file_content_type(latest_rfq_file: dict[str, str]) -> str:
    explicit_content_type = _stringify_value(latest_rfq_file.get("content_type")).lower()
    if explicit_content_type:
        return explicit_content_type
    return _guess_content_type(
        latest_rfq_file.get("name")
        or latest_rfq_file.get("download_url")
        or latest_rfq_file.get("url")
        or latest_rfq_file.get("path")
    )


def _get_rfq_file_extension(latest_rfq_file: dict[str, str]) -> str:
    source = (
        latest_rfq_file.get("name")
        or latest_rfq_file.get("download_url")
        or latest_rfq_file.get("url")
        or latest_rfq_file.get("path")
        or ""
    )
    return Path(str(source)).suffix.lower()


def _guess_content_type(value: str | None) -> str:
    guessed_type = mimetypes.guess_type(str(value or ""))[0]
    return str(guessed_type or "").lower()


def _normalize_file_role(file_entry: dict[str, str]) -> str:
    return _stringify_value(file_entry.get("file_role")).strip().upper()


def _is_pdf_file(file_entry: dict[str, str]) -> bool:
    return (
        _detect_rfq_file_content_type(file_entry) == "application/pdf"
        or _get_rfq_file_extension(file_entry) == ".pdf"
    )


def _is_image_file(file_entry: dict[str, str]) -> bool:
    content_type = _detect_rfq_file_content_type(file_entry)
    return content_type.startswith("image/") or _get_rfq_file_extension(file_entry) in DOCX_IMAGE_EXTENSIONS


def _is_previewable_file(file_entry: dict[str, str]) -> bool:
    return _is_pdf_file(file_entry) or _is_image_file(file_entry)


def _has_drawing_name_hint(file_entry: dict[str, str]) -> bool:
    normalized_name = _stringify_value(file_entry.get("name")).strip().lower()
    return any(hint in normalized_name for hint in DRAWING_NAME_HINTS)


def _has_reference_picture_name_hint(file_entry: dict[str, str]) -> bool:
    normalized_name = _stringify_value(file_entry.get("name")).strip().lower()
    return any(hint in normalized_name for hint in REFERENCE_PICTURE_NAME_HINTS)


def _has_any_name_hint(file_entry: dict[str, str], hints: tuple[str, ...] | set[str]) -> bool:
    normalized_name = _stringify_value(file_entry.get("name")).strip().lower()
    return any(hint in normalized_name for hint in hints)


def _is_reference_picture_before_file(file_entry: dict[str, str]) -> bool:
    file_role = _normalize_file_role(file_entry)
    return (
        file_role in REFERENCE_PICTURE_BEFORE_ROLE_HINTS
        or _has_any_name_hint(file_entry, REFERENCE_PICTURE_BEFORE_NAME_HINTS)
    )


def _is_reference_picture_after_file(file_entry: dict[str, str]) -> bool:
    file_role = _normalize_file_role(file_entry)
    return (
        file_role in REFERENCE_PICTURE_AFTER_ROLE_HINTS
        or _has_any_name_hint(file_entry, REFERENCE_PICTURE_AFTER_NAME_HINTS)
    )


def _is_same_rfq_file(left: dict[str, str], right: dict[str, str]) -> bool:
    if not left or not right:
        return False

    comparison_keys = ("id", "blob_name", "download_url", "url", "path", "name")
    for key in comparison_keys:
        left_value = _stringify_value(left.get(key))
        right_value = _stringify_value(right.get(key))
        if left_value and right_value and left_value == right_value:
            return True

    return False


def _get_offer_preparation_data(rfq_data: dict[str, Any] | None) -> dict[str, Any]:
    raw_value = (rfq_data or {}).get(OFFER_PREPARATION_DATA_KEY)
    return dict(raw_value) if isinstance(raw_value, dict) else {}


def _build_effective_offer_template_data(
    rfq_data: dict[str, Any] | None,
    offer_data: dict[str, Any] | None = None,
) -> dict[str, Any]:
    base_data = dict(rfq_data or {})
    offer_overrides = _get_offer_preparation_data(base_data)
    effective_data = dict(base_data)
    effective_data.update(offer_overrides)
    if isinstance(offer_data, dict):
        effective_data.update(offer_data)
    return effective_data


def _extract_docx_paragraphs(docx_bytes: bytes) -> list[str]:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zip_file:
        document_xml = zip_file.read("word/document.xml")

    root = ET.fromstring(document_xml)
    paragraphs: list[str] = []
    for paragraph in [node for node in root.iter() if node.tag == WORD_PARAGRAPH_TAG]:
        text_fragments = [
            node.text or "" for node in paragraph.iter() if node.tag == WORD_TEXT_TAG
        ]
        text = _normalize_preview_text("".join(text_fragments))
        if text:
            paragraphs.append(text)
    return paragraphs


def _render_preview_paragraph(text: str) -> str:
    escaped = html.escape(text)
    if re.fullmatch(r"\d+\)\s+.+", text):
        return f"<h2>{escaped}</h2>"
    if text.startswith("SUBJECT :"):
        return f"<h3>{escaped}</h3>"
    if text in {"Product picture for reference", "General conditions"}:
        return f"<p class='label'>{escaped}</p>"
    if text.startswith("Dear ") or text.startswith("Best Regards"):
        return f"<p class='highlight'>{escaped}</p>"
    if text.endswith(":") and len(text) <= 32:
        return f"<p class='label'>{escaped}</p>"
    if text.startswith("Offer validity:") or text.startswith(
        "AVO Carbon Group standard selling conditions apply."
    ):
        return f"<p class='muted'>{escaped}</p>"
    return f"<p>{escaped}</p>"


def _normalize_preview_text(value: str) -> str:
    normalized = value.replace("\u00a0", " ")
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def _format_target_price_display(data: dict[str, Any]) -> str:
    eur_value = _pick_first_raw_value(data, ("target_price_eur", "targetPrice"))
    local_value = _pick_first_raw_value(data, ("target_price_local", "targetPriceLocal"))
    local_currency = _pick_first_raw_value(
        data,
        ("target_price_currency", "targetPriceCurrency"),
    )

    eur_text = _stringify_value(eur_value).strip() if eur_value is not None else ""
    local_text = _stringify_value(local_value).strip() if local_value is not None else ""
    currency_text = (
        _stringify_value(local_currency).strip().upper()
        if local_currency is not None
        else ""
    )

    if eur_text and local_text and currency_text and currency_text != "EUR":
        return f"{eur_text} EUR / {local_text} {currency_text}".strip()
    if eur_text:
        return f"{eur_text} EUR"
    if local_text:
        return f"{local_text} {currency_text}".strip()
    return ""


def _pick_first_raw_value(data: dict[str, Any], keys: tuple[str, ...]) -> Any | None:
    for key in keys:
        if key not in data:
            continue
        value = data.get(key)
        if _has_meaningful_value(value):
            return value
    return None


def _pick_first_value(data: dict[str, Any], keys: tuple[str, ...]) -> str:
    value = _pick_first_raw_value(data, keys)
    return _stringify_value(value)


def _has_meaningful_value(value: Any) -> bool:
    if value is None:
        return False
    if isinstance(value, str):
        return bool(value.strip())
    if isinstance(value, (list, tuple, set, dict)):
        return bool(value)
    return True


def _stringify_value(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    if isinstance(value, datetime):
        dt_value = value.astimezone(timezone.utc) if value.tzinfo else value.replace(
            tzinfo=timezone.utc
        )
        return dt_value.strftime("%Y-%m-%d")
    if isinstance(value, (int, float)):
        return str(value)
    if isinstance(value, str):
        return value.strip()
    if isinstance(value, dict):
        parts = [
            f"{key}: {_stringify_value(item)}"
            for key, item in value.items()
            if _has_meaningful_value(item)
        ]
        return ", ".join(parts)
    if isinstance(value, (list, tuple, set)):
        parts = [_stringify_value(item) for item in value if _has_meaningful_value(item)]
        return ", ".join(part for part in parts if part)
    return str(value)
