from __future__ import annotations

import base64
import io
import zipfile
from datetime import datetime
from xml.etree import ElementTree as ET

from docx import Document
from reportlab.pdfgen import canvas

from app.models.rfq import Rfq, RfqPhase, RfqSubStatus
from app.services.offer_template import (
    render_offer_preparation_docx,
    render_offer_preparation_preview_html,
)


WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
SAMPLE_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAusB9sZS8n8AAAAASUVORK5CYII="
)


def _extract_document_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes), "r") as zip_file:
        return zip_file.read("word/document.xml").decode("utf-8")


def _extract_visible_document_text(docx_bytes: bytes) -> str:
    root = ET.fromstring(_extract_document_xml(docx_bytes))
    text_parts = []
    for node in root.iter(f"{{{WORD_NAMESPACE}}}t"):
        text_parts.append(node.text or "")
    return "".join(text_parts)


def _extract_paragraph_text(docx_bytes: bytes, paragraph_index: int) -> str:
    root = ET.fromstring(_extract_document_xml(docx_bytes))
    paragraphs = list(root.iter(f"{{{WORD_NAMESPACE}}}p"))
    paragraph = paragraphs[paragraph_index]
    parts: list[str] = []
    for node in paragraph.iter():
        if node.tag == f"{{{WORD_NAMESPACE}}}t":
            parts.append(node.text or "")
        elif node.tag == f"{{{WORD_NAMESPACE}}}tab":
            parts.append("\t")
        elif node.tag == f"{{{WORD_NAMESPACE}}}br":
            parts.append("\n")
    return "".join(parts)


def _extract_all_paragraph_texts(docx_bytes: bytes) -> list[str]:
    root = ET.fromstring(_extract_document_xml(docx_bytes))
    paragraphs = list(root.iter(f"{{{WORD_NAMESPACE}}}p"))
    return [_extract_paragraph_text(docx_bytes, index) for index in range(len(paragraphs))]


def _paragraph_contains_image(paragraph) -> bool:
    return "graphicData" in paragraph._p.xml


def test_render_offer_preparation_docx_replaces_detected_template_fields():
    rfq = Rfq(
        rfq_id="rfq-offer-001",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 9, 30, 0),
        rfq_data={
            "systematic_rfq_id": "26001-ASS-00",
            "customer_name": "Acme Mobility",
            "contact_name": "Jane Buyer",
            "contact_phone": "+216 11 222 333",
            "contact_email": "jane.buyer@example.com",
            "product_name": "Carbon Brush",
            "project_name": "Phoenix",
            "customer_pn": "PN-7788",
            "revision_level": "B2",
            "sop_year": "2028",
            "annual_volume": "450000",
            "type_of_packaging": "Returnable trays",
            "expected_delivery_conditions": "DAP Tunis",
            "expected_payment_terms": "60 days end of month",
            "target_price_eur": "2.45",
            "rfq_files": [
                {
                    "name": "phoenix-drawing.pdf",
                    "uploaded_at": "2026-05-05T10:00:00+00:00",
                }
            ],
        },
    )

    creator_profile = {
        "created_by_name": "Sam Seller",
        "created_by_phone": "+216 55 444 333",
        "created_by_email": "seller@example.com",
    }

    document_bytes = render_offer_preparation_docx(rfq, creator_profile=creator_profile)
    document_xml = _extract_document_xml(document_bytes)
    visible_text = _extract_visible_document_text(document_bytes)

    assert "{{customer_name}}" not in document_xml
    assert "Tuesday, May 5, 2026" in visible_text
    assert "Our reference: 20260505-Phoenix-Carbon Brush-Acme Mobility" in visible_text
    assert "RFQ file: phoenix-drawing.pdf" in visible_text
    assert "Acme Mobility" in visible_text
    assert "Sam Seller" in visible_text
    assert "+216 55 444 333" in visible_text
    assert "seller@example.com" in visible_text
    assert "Jane Buyer" in visible_text
    assert "jane.buyer@example.com" not in visible_text
    assert "Returnable trays" in visible_text
    assert "DAP Tunis" in visible_text


def test_render_offer_preparation_preview_html_contains_filled_content():
    rfq = Rfq(
        rfq_id="rfq-offer-002",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 14, 0, 0),
        rfq_data={
            "systematic_rfq_id": "26002-BRU-01",
            "customer_name": "Nova Drives",
            "contact_name": "Alex Carter",
            "contact_phone": "+33 123456789",
            "contact_email": "alex.carter@example.com",
            "product_name": "Motor Brush",
            "project_name": "Falcon",
            "customer_pn": "MD-123",
            "revision_level": "07",
            "sop_year": "2027",
            "annual_volume": "120000",
            "type_of_packaging": "Boxes",
            "expected_delivery_conditions": "FCA Plant",
            "expected_payment_terms": "45 days",
            "target_price_eur": "1.90",
            "rfq_files": [
                {
                    "name": "falcon-drawing.pdf",
                    "uploaded_at": "2026-05-05T14:15:00+00:00",
                }
            ],
        },
    )

    creator_profile = {
        "created_by_name": "Nadia Sales",
        "created_by_phone": "+33 111222333",
        "created_by_email": "seller@example.com",
    }

    preview_html = render_offer_preparation_preview_html(
        rfq,
        creator_profile=creator_profile,
    )

    assert "Offer preparation preview" in preview_html
    assert 'class="paper"' in preview_html
    assert "Tuesday, May 5, 2026" in preview_html
    assert "Our reference: 20260505-Falcon-Motor Brush-Nova Drives" in preview_html
    assert "RFQ file: falcon-drawing.pdf" in preview_html
    assert "Nova Drives" in preview_html
    assert "Alex Carter" in preview_html
    assert "Nadia Sales" in preview_html
    assert "seller@example.com" in preview_html
    assert "FCA Plant" in preview_html
    assert "45 days" in preview_html


def test_render_offer_preparation_docx_preserves_spaces_and_line_breaks():
    rfq = Rfq(
        rfq_id="rfq-offer-003",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 11, 15, 0),
        rfq_data={
            "customer_name": "nidec",
            "contact_name": "Taha Khiari",
            "contact_phone": "+216 111",
            "contact_email": "taha@example.com",
            "product_name": "Busbar",
            "project_name": "nid-001",
            "customer_pn": "p125425",
            "revision_level": "01",
            "sop_year": "2027",
            "annual_volume": "12000",
            "type_of_packaging": "Boxes",
            "expected_delivery_conditions": "FCA",
            "expected_payment_terms": "45 days",
            "target_price_eur": "4.51775",
            "target_price_local": "500",
            "target_price_currency": "INR",
            "rfq_files": [
                {
                    "name": "busbar-layout.pdf",
                    "uploaded_at": "2026-05-05T12:00:00+00:00",
                }
            ],
        },
    )

    document_bytes = render_offer_preparation_docx(rfq)

    assert _extract_paragraph_text(document_bytes, 13) == "Dear nidec,"
    assert (
        _extract_paragraph_text(document_bytes, 14)
        == "Thank you for your RFQ. We are pleased to submit our commercial offer for the supply of Busbar nid-001 manufactured according to your drawing p125425 01 (see Annex 1 and reference picture below)."
    )
    assert (
        _extract_paragraph_text(document_bytes, 19)
        == "Product: Busbar nid-001 – Drawing p125425 01\nSOP: 2027\nValidation batch: \nSerial volume assumption: 12000 pcs/year"
    )
    assert _extract_paragraph_text(document_bytes, 27) == "Quantity: 12000 pcs/year"
    assert "RFQ file: busbar-layout.pdf" in _extract_all_paragraph_texts(document_bytes)


def test_render_offer_preparation_docx_embeds_pdf_preview_when_source_file_exists(tmp_path):
    drawing_path = tmp_path / "drawing-preview.pdf"
    pdf_canvas = canvas.Canvas(str(drawing_path))
    pdf_canvas.drawString(72, 720, "Drawing preview for annex 1")
    pdf_canvas.rect(72, 520, 300, 120)
    pdf_canvas.save()

    rfq = Rfq(
        rfq_id="rfq-offer-004",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 12, 0, 0),
        rfq_data={
            "customer_name": "SAGEMCOM",
            "contact_name": "Drawing Contact",
            "product_name": "wire harness",
            "project_name": "GEN2",
            "customer_pn": "11D010301100",
            "revision_level": "A1",
            "rfq_files": [
                {
                    "name": "drawing-preview.pdf",
                    "path": str(drawing_path),
                    "content_type": "application/pdf",
                    "uploaded_at": "2026-05-05T12:15:00+00:00",
                }
            ],
        },
    )

    document_bytes = render_offer_preparation_docx(rfq)
    visible_text = _extract_visible_document_text(document_bytes)

    with zipfile.ZipFile(io.BytesIO(document_bytes), "r") as zip_file:
        media_files = [
            name for name in zip_file.namelist() if name.startswith("word/media/")
        ]

    assert "RFQ file: drawing-preview.pdf" not in visible_text
    assert media_files


def test_render_offer_preparation_docx_embeds_reference_pictures_under_template_heading(tmp_path):
    drawing_path = tmp_path / "drawing-preview.pdf"
    reference_picture_path = tmp_path / "reference-picture-1.png"
    second_reference_picture_path = tmp_path / "product-photo-2.png"

    pdf_canvas = canvas.Canvas(str(drawing_path))
    pdf_canvas.drawString(72, 720, "Drawing preview for annex 1")
    pdf_canvas.rect(72, 520, 300, 120)
    pdf_canvas.save()

    reference_picture_path.write_bytes(SAMPLE_PNG_BYTES)
    second_reference_picture_path.write_bytes(SAMPLE_PNG_BYTES)

    rfq = Rfq(
        rfq_id="rfq-offer-004b",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 12, 30, 0),
        rfq_data={
            "customer_name": "SAGEMCOM",
            "contact_name": "Drawing Contact",
            "product_name": "wire harness",
            "project_name": "GEN2",
            "customer_pn": "11D010301100",
            "revision_level": "A1",
            "rfq_files": [
                {
                    "name": "drawing-preview.pdf",
                    "path": str(drawing_path),
                    "content_type": "application/pdf",
                    "uploaded_at": "2026-05-05T12:15:00+00:00",
                },
                {
                    "name": "reference-picture-1.png",
                    "path": str(reference_picture_path),
                    "content_type": "image/png",
                    "uploaded_at": "2026-05-05T12:20:00+00:00",
                },
                {
                    "name": "product-photo-2.png",
                    "path": str(second_reference_picture_path),
                    "content_type": "image/png",
                    "uploaded_at": "2026-05-05T12:25:00+00:00",
                },
            ],
        },
    )

    document_bytes = render_offer_preparation_docx(rfq)
    document = Document(io.BytesIO(document_bytes))
    heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "Product picture for reference"
    )

    assert len(document.inline_shapes) >= 3
    assert heading_index > 0
    assert _paragraph_contains_image(document.paragraphs[heading_index - 1])


def test_render_offer_preparation_docx_supports_reference_pictures_before_and_after_heading(tmp_path):
    drawing_path = tmp_path / "drawing-preview.pdf"
    before_reference_picture_path = tmp_path / "before-product-picture-1.png"
    after_reference_picture_path = tmp_path / "after-product-picture-1.png"

    pdf_canvas = canvas.Canvas(str(drawing_path))
    pdf_canvas.drawString(72, 720, "Drawing preview for annex 1")
    pdf_canvas.rect(72, 520, 300, 120)
    pdf_canvas.save()

    before_reference_picture_path.write_bytes(SAMPLE_PNG_BYTES)
    after_reference_picture_path.write_bytes(SAMPLE_PNG_BYTES)

    rfq = Rfq(
        rfq_id="rfq-offer-004c",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 12, 45, 0),
        rfq_data={
            "customer_name": "SAGEMCOM",
            "contact_name": "Drawing Contact",
            "product_name": "wire harness",
            "project_name": "GEN2",
            "customer_pn": "11D010301100",
            "revision_level": "A1",
            "rfq_files": [
                {
                    "name": "drawing-preview.pdf",
                    "path": str(drawing_path),
                    "content_type": "application/pdf",
                    "uploaded_at": "2026-05-05T12:15:00+00:00",
                },
                {
                    "name": "before-product-picture-1.png",
                    "path": str(before_reference_picture_path),
                    "content_type": "image/png",
                    "uploaded_at": "2026-05-05T12:20:00+00:00",
                },
                {
                    "name": "after-product-picture-1.png",
                    "path": str(after_reference_picture_path),
                    "content_type": "image/png",
                    "uploaded_at": "2026-05-05T12:25:00+00:00",
                },
            ],
        },
    )

    document_bytes = render_offer_preparation_docx(rfq)
    document = Document(io.BytesIO(document_bytes))
    heading_index = next(
        index
        for index, paragraph in enumerate(document.paragraphs)
        if paragraph.text.strip() == "Product picture for reference"
    )

    assert heading_index > 0
    assert _paragraph_contains_image(document.paragraphs[heading_index - 1])
    assert _paragraph_contains_image(document.paragraphs[heading_index + 1])


def test_render_offer_preparation_docx_prefers_offer_preparation_overrides():
    rfq = Rfq(
        rfq_id="rfq-offer-005",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 12, 0, 0),
        rfq_data={
            "customer_name": "Base Customer",
            "product_name": "Base Product",
            "project_name": "Base Project",
            "customer_pn": "BASE-001",
            "revision_level": "A0",
            "offer_preparation_data": {
                "customer_name": "Offer Customer",
                "product_name": "Offer Product",
                "project_name": "Offer Project",
                "customer_pn": "OFFER-001",
                "revision_level": "B1",
            },
        },
    )

    document_bytes = render_offer_preparation_docx(rfq)
    visible_text = _extract_visible_document_text(document_bytes)

    assert "Offer Customer" in visible_text
    assert "Offer Product" in visible_text
    assert "Offer Project" in visible_text
    assert "OFFER-001" in visible_text
    assert "B1" in visible_text
    assert "Base Customer" not in visible_text


def test_render_offer_preparation_docx_prefers_explicit_offer_table_data():
    rfq = Rfq(
        rfq_id="rfq-offer-005b",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 12, 30, 0),
        rfq_data={
            "customer_name": "Base Customer",
            "product_name": "Base Product",
            "project_name": "Base Project",
            "customer_pn": "BASE-001",
            "revision_level": "A0",
        },
    )

    document_bytes = render_offer_preparation_docx(
        rfq,
        offer_data={
            "customer_name": "Offer Table Customer",
            "product_name": "Offer Table Product",
            "project_name": "Offer Table Project",
            "customer_pn": "TABLE-001",
            "revision_level": "C4",
        },
    )
    visible_text = _extract_visible_document_text(document_bytes)

    assert "Offer Table Customer" in visible_text
    assert "Offer Table Product" in visible_text
    assert "Offer Table Project" in visible_text
    assert "TABLE-001" in visible_text
    assert "C4" in visible_text
    assert "Base Customer" not in visible_text


def test_render_offer_preparation_docx_fills_offer_specific_template_fields():
    rfq = Rfq(
        rfq_id="rfq-offer-006",
        phase=RfqPhase.OFFER,
        sub_status=RfqSubStatus.PREPARATION,
        created_by_email="seller@example.com",
        created_at=datetime(2026, 5, 5, 16, 45, 0),
        rfq_data={
            "customer_name": "SAGEMCOM",
            "contact_name": "Aya Contact",
            "product_name": "wire harness",
            "project_name": "GEN2",
            "customer_pn": "11D010301100",
            "revision_level": "A1",
            "sop_year": "2027",
            "annual_volume": "85000",
            "type_of_packaging": "Box with separator",
            "expected_delivery_conditions": "DAP Tunis",
            "expected_payment_terms": "60 days",
            "offer_preparation_data": {
                "your_reference": "RFQ/WATER NAM GEN2",
                "copies": "Mohamed Laith Ben Mabrouk <mohamed@example.com>",
                "subject": "commercial offer AVOCarbon TUNISIA-SAGEMCOM",
                "validation_batch": "300 pcs for customer validation",
                "pilot_quantity": "300 pcs",
                "pilot_unit_price": "3.85 EUR",
                "material_balance_moq": "1,500 pcs",
                "serial_unit_price": "2.95 EUR",
                "lead_time_deliveries": "6 weeks after PO, weekly call-off possible",
                "inventory_commitment": "Customer to release a 3-month firm forecast",
                "offer_validity": "3 months from date of issue.",
            },
        },
    )

    document_bytes = render_offer_preparation_docx(rfq)

    assert _extract_paragraph_text(document_bytes, 7) == "Your reference: RFQ/WATER NAM GEN2"
    assert (
        _extract_paragraph_text(document_bytes, 9)
        == "Copies : Mohamed Laith Ben Mabrouk <mohamed@example.com>"
    )
    assert (
        _extract_paragraph_text(document_bytes, 11)
        == "SUBJECT : commercial offer AVOCarbon TUNISIA-SAGEMCOM"
    )
    expected_scope_paragraph = (
        "Product: wire harness GEN2 "
        f"{chr(8211)} Drawing 11D010301100 A1\n"
        "SOP: 2027\n"
        "Validation batch: 300 pcs for customer validation\n"
        "Serial volume assumption: 85000 pcs/year"
    )
    assert _extract_paragraph_text(document_bytes, 19) == expected_scope_paragraph
    assert _extract_paragraph_text(document_bytes, 23) == "Quantity: 300 pcs"
    assert _extract_paragraph_text(document_bytes, 24) == "Unit price: 3.85 EUR"
    assert _extract_paragraph_text(document_bytes, 25) == "Material balance (MOQ): 1,500 pcs"
    assert _extract_paragraph_text(document_bytes, 28) == "Unit price: 2.95 EUR"
    paragraph_texts = _extract_all_paragraph_texts(document_bytes)
    assert "6 weeks after PO, weekly call-off possible" in paragraph_texts
    assert "Customer to release a 3-month firm forecast" in paragraph_texts
    assert "Offer validity: 3 months from date of issue." in paragraph_texts
